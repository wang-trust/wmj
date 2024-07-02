#coding=utf-8
import os
import sys
from PyQt5.QtWidgets import QMainWindow, QApplication, QFileDialog
from ui import Ui_MainWindow
import time
import re
import docx
from win32com import client
import openpyxl
import xlrd



class Example(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.show()

        self.ctrlConnect()
        self.defVar()

    # UI contorl function --------------------------------
    def defVar(self):
        self.fileCheckPath = ''
        self.fileWritePath = ''
        self.fileCheckArray = []
        self.fileWriteArray = []
        # 1: 公开   2：内部  4: 秘密及以上
        self.checkLevel = 0x1
        self.writeLevel = 0x1
        # [, num, ] 0:INFO 1:WARING 2:ERROR 3:FATAL
        self.resultArray = []

        # for re.sub replace other symbols
        self.resymbols2 = '[\t\n\- \(\)（）\[\]【】]+'
        # self.resymbols2Text = '[\t\n- ]+'

        # for debug
        self.log = open('run.log', 'a', encoding='utf-8')
        # self.runCheckPath = 'C:\\Users\\goupi\\Desktop\\qt5\\dist\\'

    def ctrlConnect(self):
        self.pushButton.clicked.connect(self.setCheckPath)
        self.pushButton_4.clicked.connect(self.setWritePath)
        self.pushButton_2.clicked.connect(self.runCheckPath)
        self.pushButton_3.clicked.connect(self.runWritePath)
        self.lineEdit.textChanged.connect(self.exchangeCheckPath)
        self.lineEdit_2.textChanged.connect(self.exchangeWritePath)
        self.comboBox.currentIndexChanged.connect(self.setCheckLevel)
        self.comboBox_3.currentIndexChanged.connect(self.setWriteLevel)

    def setCheckPath(self):
        file_path = QFileDialog.getExistingDirectory(self, "选择文件夹", "./", QFileDialog.ShowDirsOnly)
        if file_path:
            # Using relative paths
            fp = os.path.split(file_path)
            if(fp[1]):
                os.chdir(fp[0])
                self.fileCheckPath = fp[1]
            else:
                self.fileCheckPath = file_path
            self.lineEdit.setText(file_path)


    def setWritePath(self):
        file_path = QFileDialog.getExistingDirectory(self, "选择文件夹", "./", QFileDialog.ShowDirsOnly)
        if file_path:
            # Using relative paths
            fp = os.path.split(file_path)
            if(fp[1]):
                os.chdir(fp[0])
                self.fileWritePath = fp[1]
            else:
                self.fileWritePath = file_path
            self.lineEdit_2.setText(file_path)

    def exchangeCheckPath(self):
        file_path = self.lineEdit.text().strip()
        if(file_path):
            self.fileCheckPath = file_path

    def exchangeWritePath(self):
        file_path = self.lineEdit_2.text().strip()
        if(file_path):
            self.fileWritePath = file_path

    def setCheckLevel(self):
        com = self.comboBox.currentText()
        if(com == '内部'):
            self.checkLevel = 0x2
        else:
            self.checkLevel = 0x1

    def setWriteLevel(self):
        com = self.comboBox_3.currentText()
        if(com == '内部'):
            self.writeLevel = 0x2
        else:
            self.writeLevel = 0x1
    # UI contorl function --------------------------------
    

    # core business------------------------------------
    # core logical for check dic
    def runCheckPath(self):
        # clear
        self.resultArray = []

        # get array
        if not self.fileCheckPath:
            self.resultArray.append([self.fileCheckPath, 2, 'File path error.'])
            self.writeLog()
            return
        self.fileCheckArray = []
        try:
            # judge self.fileCheckPath is right ?
            judgetemp = os.listdir(self.fileCheckPath)
            self.recursivePath(self.fileCheckPath, self.fileCheckArray)
            # self.printArray(self.fileCheckArray)
        except Exception as e:
            self.resultArray.append([self.fileCheckPath, 2, e])
            self.writeLog()
            return

        # core function
        for each in self.fileCheckArray:
            tempfile = os.path.normpath(each)
            self.checkFile(tempfile)
        # temp = os.path.normpath(self.fileCheckArray[0])
        # filestr = 'C:\\Users\\goupi\\Desktop\\electron\\demo11\\dist\\1_().rar.part1'
        # self.checkFile(filestr)

        # write log
        self.writeLog()

    def checkFile(self, filepath):
        # clear other symbols
        titleLevel = self.checkLevel
        nameLevel = 0
        contentLevel = -1
        nameTrueLevel = 0
        contentTrueLevel = 0

        newfilepath = re.sub(self.resymbols2, '', filepath)
        if(filepath.find('秘密') != -1 or newfilepath.find('绝密') != -1 or \
            newfilepath.find('机密') != -1):
            self.resultArray.append([filepath, 3, '文件名出现秘密字眼'])
            return
        # 获取文件允许的最高密级
        filenamelist = newfilepath.split('\\')
        for each in filenamelist[:-1]:
            if(each.find('内部') != -1):
                if(titleLevel < 2):
                    self.resultArray.append([filepath, 2, '公开密级中存在内部文件'])
                    return
            elif(each.find('公开') != -1):
                if(titleLevel > 1):
                    titleLevel = 1

        # 获取文件的密级设定
        suffix = os.path.splitext(filenamelist[-1])
        if(suffix[0].find('内部') != -1):
            nameLevel |= 2
        if(suffix[0].find('公开') != -1):
            nameLevel |= 1
        if(suffix[0].find('内部') == 0 or suffix[0].endswith('内部')):
            nameTrueLevel = 2
        elif(suffix[0].find('公开') == 0 or suffix[0].endswith('公开')):
            nameTrueLevel = 1
        
        # read text-file
        if(suffix[1] == '.txt'):
            contentLevel, contentTrueLevel = self.checkTXT(filepath)
        elif(suffix[1] == '.docx'):
            contentLevel, contentTrueLevel = self.checkDOCX(filepath)
        elif(suffix[1] == '.doc'):
            contentLevel, contentTrueLevel = self.checkDOC(filepath)
        elif(suffix[1] == '.xlsx'):
            contentLevel, contentTrueLevel = self.checkXLSX(filepath)
        elif(suffix[1] == '.xls'):
            contentLevel, contentTrueLevel = self.checkXLS(filepath)
        elif(suffix[1] == '.pdf'):
            self.resultArray.append([filepath, 2, 'pdf文件需要手动确认'])
        elif(suffix[1] == '.exe'):
            self.resultArray.append([filepath, 2, '存在exe文件'])
        elif(suffix[1] == '.csv'):
            contentLevel, contentTrueLevel = self.checkTXT(filepath)
        else:
            contentLevel = -1

        # make result
        self.judgeResult(filepath, titleLevel, nameLevel, contentLevel, nameTrueLevel, contentTrueLevel)
    

    def judgeResult(self, filepath, titleLevel, nameLevel, contentLevel, nameTrueLevel, contentTrueLevel):
        if(contentLevel >= 4):
            self.resultArray.append([filepath, 3, '文件内部出现秘密字眼'])
        elif(contentLevel == 3):
            self.resultArray.append([filepath, 2, '文件内部同时出现内部和公开'])
        elif(nameLevel == 3):
            self.resultArray.append([filepath, 2, '文件名中同时出现内部和公开'])
        elif(contentLevel == 2):
            if(nameLevel == 0):
                if(titleLevel == 2):
                    self.resultArray.append([filepath, 1, '文件名和文件内仅一处标密，需打开文件确认'])
                else:
                    self.resultArray.append([filepath, 2, '公开文件中出现内部字眼'])
            elif(nameLevel == 1):
                self.resultArray.append([filepath, 2, '公开文件中出现内部字眼'])
            elif(nameLevel == 2):
                if(titleLevel == 2):
                    if(nameTrueLevel == 2 and contentTrueLevel == 2):
                        self.resultArray.append([filepath, 0, 'OK'])
                    else:
                        self.resultArray.append([filepath, 2, '密级标注错误'])
                else:
                    self.resultArray.append([filepath, 2, '公开密级中出现内部文件'])
        elif(contentLevel == 1):
            if(nameLevel == 0):
                self.resultArray.append([filepath, 1, '文件名和文件内仅一处标密，需打开文件确认'])
            elif(nameLevel == 1):
                if(nameTrueLevel > 0 and contentTrueLevel > 0):
                    self.resultArray.append([filepath, 0, 'OK'])
                else:
                    self.resultArray.append([filepath, 2, '密级标注错误'])
            elif(nameLevel == 2):
                if(titleLevel == 2):
                    self.resultArray.append([filepath, 2, '内部文件中出现公开字眼'])
                else:
                    self.resultArray.append([filepath, 2, '公开密级中出现内部文件'])
        elif(contentLevel == 0):
            if(nameLevel == 0):
                self.resultArray.append([filepath, 2, '文件未标密'])
            elif(nameLevel == 1):
                self.resultArray.append([filepath, 1, '文件名和文件内仅一处标密，需打开文件确认'])
            elif(nameLevel == 2):
                self.resultArray.append([filepath, 1, '文件名和文件内仅一处标密，需打开文件确认'])
        elif(contentLevel == -1):
            if(nameLevel == 0):
                self.resultArray.append([filepath, 2, '文件未标密'])
            elif(nameLevel == 1):
                if(nameTrueLevel > 0):
                    self.resultArray.append([filepath, 0, 'OK'])
                else:
                    self.resultArray.append([filepath, 2, '密级标注错误'])
            elif(nameLevel == 2):
                if(titleLevel == 2):
                    if(nameTrueLevel == 2):
                        self.resultArray.append([filepath, 0, 'OK'])
                    else:
                        self.resultArray.append([filepath, 2, '密级标注错误'])
                else:
                    self.resultArray.append([filepath, 2, '公开密级中出现内部文件'])
        else:
            self.resultArray.append([filepath, 2, '未知错误'])

    def checkXLS(self, filepath):
        resultLevel = 0
        resultTrueLevel = 0

        try:
            wb = xlrd.open_workbook(filepath)
            sheetnames = wb.sheet_names()
        except Exception as e:
            self.resultArray.append([filepath, 2, e])
            return resultLevel, resultTrueLevel
        for sheetname in sheetnames:
            tmp = re.sub(self.resymbols2, '', sheetname)
            if(tmp.find('秘密') != -1 or \
                tmp.find('机密') != -1 or \
                tmp.find('绝密') != -1):
                resultLevel |= 0x04
                return resultLevel
            if(tmp.find('内部') != -1):
                resultLevel |= 0x02
            if(tmp.find('公开') != -1):
                resultLevel |= 0x01
        for sheetname in sheetnames:
            sheet = wb.sheet_by_name(sheetname)
            for i in range(sheet.nrows):
                for j in range(sheet.ncols):
                    v1 = sheet.cell(i, j).value
                    if(type(v1) == str):
                        sheetText = re.sub(self.resymbols2, '', v1)
                        if(sheetText.find('秘密') != -1 or \
                            sheetText.find('机密') != -1 or \
                            sheetText.find('绝密') != -1):
                            resultLevel |= 0x04
                        if(sheetText.find('内部') != -1):
                            resultLevel |= 0x02
                        if(sheetText.find('公开') != -1):
                            resultLevel |= 0x01
        firstsheet = wb.sheet_by_name(sheetnames[0])
        for i in range(firstsheet.nrows):
            for j in range(firstsheet.ncols):
                v1 = firstsheet.cell(i, j).value
                if(type(v1) == str):
                    sheetText = re.sub(self.resymbols2, '', firstsheet.cell(i, j).value)
                    if(sheetText.find('内部') == 0):
                        resultTrueLevel = 2
                        return resultLevel, resultTrueLevel
                    elif(sheetText.find('公开') == 0):
                        resultTrueLevel = 1
                        return resultLevel, resultTrueLevel

        return resultLevel, resultTrueLevel

    def checkXLSX(self, filepath):
        # workbook = openpyxl.load_workbook(filepath + '123.xlsx')
        resultLevel = 0
        resultTrueLevel = 0
        try:
            wb = openpyxl.load_workbook(filepath)
        except Exception as e:
            self.resultArray.append([filepath, 2, e])
            return resultLevel, resultTrueLevel

        # 获取各个工作表
        for sheet in wb:
            tmp = re.sub(self.resymbols2, '', sheet.title)
            if(tmp.find('秘密') != -1 or \
                tmp.find('机密') != -1 or \
                tmp.find('绝密') != -1):
                resultLevel |= 0x04
                return resultLevel
            if(tmp.find('内部') != -1):
                resultLevel |= 0x02
            if(tmp.find('公开') != -1):
                resultLevel |= 0x01
        # for 遍历分析
        # column 列  row 行
        for sheet in wb:
            for i in range(1, sheet.max_column + 1):
                for j in range(1, sheet.max_row + 1):
                    # print('[{0},{1}]={2}'.format(i, j, sheet.cell(i, j).value))
                    v1 = sheet.cell(i, j).value
                    # print(type(sheet.cell(i, j).value))
                    if(type(v1) == str):
                        sheetText = re.sub(self.resymbols2, '', v1)
                        if(sheetText.find('秘密') != -1 or \
                            sheetText.find('机密') != -1 or \
                            sheetText.find('绝密') != -1):
                            resultLevel |= 0x04
                        if(sheetText.find('内部') != -1):
                            resultLevel |= 0x02
                        if(sheetText.find('公开') != -1):
                            resultLevel |= 0x01
        # for 公开
        for sheet in wb:
            for i in range(1, sheet.max_column + 1):
                for j in range(1, sheet.max_row + 1):
                    v1 = sheet.cell(i, j).value
                    if(type(v1) == str):
                        sheetText = re.sub(self.resymbols2, '', sheet.cell(i, j).value)
                        if(sheetText.find('内部') == 0):
                            resultTrueLevel = 2
                            return resultLevel, resultTrueLevel
                        elif(sheetText.find('公开') == 0):
                            resultTrueLevel = 1
                            return resultLevel, resultTrueLevel
            break
        return resultLevel, resultTrueLevel

    def checkDOC(self, filepath):
        resultLevel = 0
        resultTrueLevel = 0
        try:
            a = os.path.split(filepath)
            b = os.path.splitext(a[-1])[0]
            newdocx = "{}\\{}----.docx".format(a[0], b)
        except Exception as e:
            self.resultArray.append([filepath, 2, e])
            return resultLevel, resultTrueLevel
        try:
            word = client.Dispatch("Word.Application")
            doc = word.Documents.Open(filepath)
            doc.SaveAs(newdocx, 12)
            doc.Close()
            word.Quit()
            resultLevel, resultTrueLevel = self.checkDOCX(newdocx, 0)
            time.sleep(1)
            os.remove(newdocx)
        except Exception as e:
            self.resultArray.append([filepath, 2, e])
        return resultLevel, resultTrueLevel

    def checkDOCX(self, filepath, flag = 1):
        contentText = ''
        resultLevel = 0
        resultTrueLevel = 0
        try:
            newdocx = docx.Document(filepath)
        except Exception as e:
            if(flag):
                self.resultArray.append([filepath, 2, e])
            return resultLevel, resultTrueLevel

        # read text
        for each in newdocx.paragraphs:
            contentText += each.text
        content = re.sub(self.resymbols2, '', contentText)
        if(content.find('秘密') != -1 \
            or content.find('机密') != -1 \
            or content.find('绝密') != -1 ):
            resultLevel |= 0x04
        if(content.find('内部') != -1):
            resultLevel |= 0x02
        if(content.find('公开') != -1):
            resultLevel |= 0x01
        # read tables
        tables = newdocx.tables
        for t in tables:
            for i in range(0, len(t.rows)):
                for j in range(0, len(t.columns)):
                    cellText = t.cell(i, j).text
                    cellText = re.sub(self.resymbols2, '', cellText)
                    if(cellText.find('秘密') != -1 \
                        or cellText.find('机密') != -1 \
                        or cellText.find('绝密') != -1 ):
                        resultLevel |= 0x04
                    if(cellText.find('内部') != -1):
                        resultLevel |= 0x02
                    if(cellText.find('公开') != -1):
                        resultLevel |= 0x01

        if(content.find('内部') == 0):
            resultTrueLevel = 2
        elif(content.find('公开') == 0):
            resultTrueLevel = 1
        if(not resultTrueLevel):
            if(len(newdocx.tables)):
                t0 = newdocx.tables[0]
                for i in range(0, len(t0.rows)):
                    for j in range(0, len(t0.columns)):
                        cellText = t0.cell(i, j).text
                        cellText = re.sub(self.resymbols2, '', cellText)
                        if(cellText == '内部'):
                            resultTrueLevel = 2
                            return resultLevel, resultTrueLevel
                        elif(cellText == '公开'):
                            resultTrueLevel = 1
                            return resultLevel, resultTrueLevel
        return resultLevel, resultTrueLevel

    def checkTXT(self, filepath):
        gbkflag = 0
        utf8flag = 0
        contentLevel = 0
        contentTrueLevel = 0
        # gbk try read txt
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                content = f.read(-1)
                content = re.sub(self.resymbols2, '', content)
            gbkflag = 1
        except Exception as e:
            # print(e)
            gbkflag = 0
            content = None
            self.resultArray.append([filepath, 0, '.txt not gbk encode'])
        # utf-8 try read txt
        if(not gbkflag):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read(-1)
                    content = re.sub(self.resymbols2, '', content)
                utf8flag = 1
            except Exception as e:
                # print(e)
                self.resultArray.append([filepath, 0, '.txt not utf-8 encode'])
                content = None
                utf8flag = 0
        if(gbkflag or utf8flag):
            if(content.find('秘密') != -1 \
                or content.find('机密') != -1 \
                or content.find('绝密') != -1 ):
                contentLevel |= 0x04
            if(content.find('内部') != -1):
                contentLevel |= 0x02
            if(content.find('公开') != -1):
                contentLevel |= 0x01
            if(content.find('内部') == 0):
                contentTrueLevel = 2
            elif(content.find('公开') == 0):
                contentTrueLevel = 1
        else:
            self.resultArray.append([filepath, 1, '.txt open fail'])
        return contentLevel, contentTrueLevel
    # core business------------------------------------


    def runWritePath(self):
        # get array
        if not self.fileWritePath:
            self.wlog(2, self.fileWritePath + 'File path error')
            return
        self.fileWriteArray = []
        try:
            # judge self.fileCheckPath is right ?
            judgetemp = os.listdir(self.fileWritePath)
            self.recursivePath(self.fileWritePath, self.fileWriteArray)
        except Exception as e:
            self.wlog(2, e)
            return
        # self.printArray(self.fileWriteArray)

        print(self.fileCheckPath)
        print(self.fileWritePath)

        # core logic
        for each in self.fileWriteArray:
            tempfile = os.path.normpath(each)
            self.writeFile(tempfile)

        # log
        

    def writeFile(self, filepath):
        print(filepath)
        return


    # helper function --------------------------------------
    def wlog(self, num, msg):
        '''0  1  2  3  4'''
        timestr = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())
        if(num == 0):
            self.log.write('INFO: {0}, {1}.\n'.format(timestr, msg))
        elif(num == 1):
            self.log.write('WARNING: {0}, {1}.\n'.format(timestr, msg))
        elif(num == 2):
            self.log.write('ERROR: {0}, {1}.\n'.format(timestr, msg))
        elif(num == 3):
            self.log.write('FATAL: {0}, {1}.\n'.format(timestr, msg))
        self.log.flush()
            
    def writeLog(self):
        infonum = 0
        warningnum = 0
        errornum = 0
        fatalnum = 0
        with open('result.log', 'w', encoding='utf-8') as f:
            timestr = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())
            f.write('Run time: {0}\n'.format(timestr))
            for each in self.resultArray:
                if(each[1] == 0):
                    f.write('INFO: {0} {1}\n'.format(os.path.abspath(each[0]), each[2]))
                    infonum += 1
                elif(each[1] == 1):
                    f.write('WARNING: {0} {1}\n'.format(os.path.abspath(each[0]), each[2]))
                    warningnum += 1
                elif(each[1] == 2):
                    f.write('ERROR: {0} {1}\n'.format(os.path.abspath(each[0]), each[2]))
                    errornum += 1
                elif(each[1] == 3):
                    f.write('FATAL: {0} {1}\n'.format(os.path.abspath(each[0]), each[2]))
                    fatalnum += 1
            
            f.write('\nResult: INFO: {0}, WARNING: {1}, ERROR: {2}, FATAL: {3}\n\n' \
                .format(infonum, warningnum, errornum, fatalnum))
        os.system('start notepad result.log')
        print('Run over...')



    def recursivePath(self, filepath, fileArray):
        if(os.path.isdir(filepath)):
            for each in os.listdir(filepath):
                each = filepath + os.sep + each
                self.recursivePath(each, fileArray)
        else:
            fileArray.append(filepath)




    # test function
    def printArray(self, l):
        for each in l:
            print(each)

    # helper function --------------------------------------



if __name__ == '__main__':
    app = QApplication(sys.argv)
    w = Example()
    sys.exit(app.exec_())
